from docx import Document
import os
import random
import requests
import json

# API configuration
API_BASE_URL = "http://localhost:8000"
ADMIN_USERNAME = "admin"  # Default admin user created by create_admin.py
ADMIN_PASSWORD = "admin123"  # Default password

# Define 5 different topics
TOPICS = [
    "Technology",
    "Science",
    "History",
    "Literature",
    "Sports"
]

# Sample content templates for each topic
CONTENT_TEMPLATES = {
    "Technology": [
        "The rapid advancement of artificial intelligence is transforming industries worldwide. Machine learning algorithms are becoming increasingly sophisticated, enabling applications in healthcare, finance, and autonomous vehicles.",
        "Cloud computing has revolutionized how businesses operate, providing scalable infrastructure and reducing costs. Companies can now deploy applications globally with unprecedented speed and reliability.",
        "The Internet of Things (IoT) connects billions of devices, creating smart homes, cities, and industrial systems. This connectivity enables real-time monitoring and automated decision-making.",
        "Blockchain technology offers decentralized solutions for secure transactions and data management. Cryptocurrencies and smart contracts are just the beginning of its potential applications."
    ],
    "Science": [
        "Quantum physics explores the fundamental nature of matter and energy at the smallest scales. Quantum entanglement and superposition challenge our classical understanding of reality.",
        "Climate change research reveals the urgent need for sustainable practices. Rising temperatures and extreme weather events are already impacting ecosystems worldwide.",
        "Genetic engineering techniques like CRISPR allow precise editing of DNA sequences. This technology holds promise for treating genetic diseases and improving crop yields.",
        "Space exploration continues to expand our understanding of the universe. Missions to Mars and beyond will require innovative solutions for long-duration space travel."
    ],
    "History": [
        "The Industrial Revolution marked a turning point in human history, shifting societies from agrarian economies to industrialized nations. Technological innovations drove unprecedented economic growth.",
        "Ancient civilizations developed complex societies with advanced mathematics, astronomy, and engineering. The legacy of these cultures continues to influence modern civilization.",
        "World War II reshaped global politics and led to the formation of international organizations. The conflict's technological advancements accelerated scientific progress.",
        "The Renaissance period witnessed a revival of art, science, and humanism. This cultural movement laid the foundation for the modern world and scientific revolution."
    ],
    "Literature": [
        "Shakespeare's works continue to captivate audiences with their timeless themes of love, betrayal, and human nature. His plays explore the complexities of the human condition.",
        "Modernist literature broke traditional narrative structures, experimenting with stream-of-consciousness and fragmented storytelling. Authors like James Joyce revolutionized literary forms.",
        "Science fiction explores alternative futures and technological possibilities. These narratives often serve as cautionary tales about the impact of technology on society.",
        "Poetry expresses emotions and ideas through rhythm, imagery, and metaphor. Contemporary poets continue to push boundaries and challenge conventional forms."
    ],
    "Sports": [
        "Olympic Games bring together athletes from around the world in a celebration of human achievement. The competition fosters international cooperation and cultural exchange.",
        "Team sports require coordination, strategy, and individual skill. Successful teams balance talent with effective communication and leadership.",
        "Athletes train rigorously to achieve peak physical performance. Sports science combines physiology, nutrition, and psychology to optimize training methods.",
        "Esports have emerged as a competitive field, attracting millions of participants worldwide. Professional gaming requires strategic thinking and rapid decision-making."
    ]
}

def authenticate():
    """Authenticate with the API and return JWT token"""
    login_url = f"{API_BASE_URL}/api/auth/login"
    login_data = {
        "username": ADMIN_USERNAME,
        "password": ADMIN_PASSWORD
    }

    try:
        response = requests.post(login_url, data=login_data)
        response.raise_for_status()
        token_data = response.json()
        return token_data["access_token"]
    except requests.exceptions.RequestException as e:
        print(f"Authentication failed: {e}")
        return None

def get_auth_headers(token):
    """Get authorization headers for API requests"""
    return {
        "Authorization": f"Bearer {token}"
    }

def clear_all_documents(token):
    """Clear all documents from the server"""
    print("Clearing all documents from server...")

    # First, get all documents
    list_url = f"{API_BASE_URL}/api/documents/"
    headers = get_auth_headers(token)

    try:
        response = requests.get(list_url, headers=headers)
        response.raise_for_status()
        documents = response.json()

        if not documents:
            print("No documents to clear.")
            return

        print(f"Found {len(documents)} documents to delete.")

        # Delete each document
        deleted_count = 0
        for doc in documents:
            delete_url = f"{API_BASE_URL}/api/documents/{doc['id']}"
            try:
                delete_response = requests.delete(delete_url, headers=headers)
                delete_response.raise_for_status()
                deleted_count += 1
                print(f"  Deleted document: {doc['original_filename']}")
            except requests.exceptions.RequestException as e:
                print(f"  Failed to delete document {doc['id']}: {e}")

        print(f"Successfully deleted {deleted_count} documents.")

    except requests.exceptions.RequestException as e:
        print(f"Failed to list documents: {e}")

def upload_document(token, file_path, description="", tags=None, is_public="private"):
    """Upload a single document to the server"""
    upload_url = f"{API_BASE_URL}/api/documents/upload"
    headers = get_auth_headers(token)

    try:
        with open(file_path, 'rb') as f:
            files = {
                'file': (os.path.basename(file_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            data = {
                'description': description,
                'is_public': is_public
            }

            if tags:
                data['tags'] = tags

            response = requests.post(upload_url, headers=headers, files=files, data=data)
            response.raise_for_status()

            result = response.json()
            print(f"  Uploaded: {os.path.basename(file_path)} (ID: {result['id']})")
            return result

    except requests.exceptions.RequestException as e:
        print(f"  Failed to upload {os.path.basename(file_path)}: {e}")
        return None

def import_generated_documents(token, output_dir):
    """Import all generated documents to the server"""
    print("Importing generated documents to server...")

    uploaded_count = 0
    topic_counts = {topic: 0 for topic in TOPICS}

    for topic in TOPICS:
        print(f"Uploading documents for topic: {topic}")

        for i in range(1, 5):  # 4 documents per topic
            filename = f"{topic.lower()}_doc_{i}.docx"
            filepath = os.path.join(output_dir, filename)

            if not os.path.exists(filepath):
                print(f"  Warning: {filename} not found, skipping.")
                continue

            # Create description and tags based on topic
            description = f"Sample document about {topic.lower()} - Document {i}"
            # tags_str = f"{topic.lower()},sample-{i}"  # Disabled due to upload issue with tags

            result = upload_document(token, filepath, description, None, "public")
            if result:
                uploaded_count += 1
                topic_counts[topic] += 1

    print(f"\nSuccessfully uploaded {uploaded_count} documents.")
    for topic, count in topic_counts.items():
        print(f"  {topic}: {count} documents")

def generate_document(topic, doc_number):
    """Generate a single DOCX document for a given topic."""
    doc = Document()

    # Add title
    title = doc.add_heading(f'{topic} Document {doc_number}', 0)

    # Add introduction paragraph
    intro = doc.add_paragraph(f"This document explores various aspects of {topic.lower()}. ")

    # Add 3-5 content paragraphs
    num_paragraphs = random.randint(3, 5)
    available_content = CONTENT_TEMPLATES[topic].copy()

    for _ in range(num_paragraphs):
        if available_content:
            content = random.choice(available_content)
            available_content.remove(content)
            doc.add_paragraph(content)
        else:
            # If we run out of templates, create generic content
            doc.add_paragraph(f"Additional insights into {topic.lower()} continue to emerge as research and development progress in this field.")

    # Add conclusion
    doc.add_paragraph(f"In conclusion, {topic.lower()} remains a dynamic and evolving field with significant implications for our future.")

    return doc

def main():
    """Generate 20 DOCX documents across 5 topics and optionally manage server documents."""
    # Create output directory if it doesn't exist
    output_dir = "test_documents"
    os.makedirs(output_dir, exist_ok=True)

    doc_count = 1

    for topic in TOPICS:
        print(f"Generating documents for topic: {topic}")

        # Generate 4 documents per topic
        for i in range(1, 5):
            doc = generate_document(topic, i)
            filename = f"{topic.lower()}_doc_{i}.docx"
            filepath = os.path.join(output_dir, filename)

            doc.save(filepath)
            print(f"  Created: {filepath}")

            doc_count += 1

    print(f"\nSuccessfully generated {doc_count - 1} documents in '{output_dir}' directory")
    print("Topics covered:", ", ".join(TOPICS))

    # Server management
    print("\n" + "="*50)
    print("SERVER MANAGEMENT")
    print("="*50)

    # Authenticate with server
    print("Authenticating with server...")
    token = authenticate()
    if not token:
        print("Authentication failed. Skipping server operations.")
        return

    print("Authentication successful.")

    # Clear existing documents
    clear_all_documents(token)

    # Import generated documents
    import_generated_documents(token, output_dir)

    print("\nDocument generation and server import completed!")

if __name__ == "__main__":
    main()